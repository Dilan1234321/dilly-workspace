"""
ATS-safe DOCX resume generator.

Generates a clean, single-column DOCX resume from the same structured
sections JSON that drives the PDF builder. DOCX is more forgiving than
PDF for most ATS parsers — it's also the format some old-school parsers
(Taleo, iCIMS-native) prefer. Formatting rules intentionally match the
PDF builder so a resume generated once can be delivered in either
format without the parser seeing different content.

Usage:
    docx_bytes = build_ats_docx(sections, ats_system="greenhouse")
"""

from __future__ import annotations

import io
import re
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


_WORKDAY_SYSTEMS = {"workday"}

_SECTION_HEADERS: dict[str, str] = {
    "education": "EDUCATION",
    "professional_experience": "EXPERIENCE",
    "research": "RESEARCH EXPERIENCE",
    "campus_involvement": "LEADERSHIP & INVOLVEMENT",
    "volunteer_experience": "VOLUNTEER EXPERIENCE",
    "projects": "PROJECTS",
    "skills": "SKILLS",
    "certifications": "CERTIFICATIONS",
    "honors": "HONORS & AWARDS",
    "coursework": "RELEVANT COURSEWORK",
    "publications": "PUBLICATIONS",
    "summary": "SUMMARY",
}

_WORKDAY_HEADERS: dict[str, str] = {
    **_SECTION_HEADERS,
    "professional_experience": "WORK EXPERIENCE",
}


def _format_date_for_ats(date_str: str, ats_system: str) -> str:
    if not date_str or not date_str.strip():
        return date_str or ""
    date_str = date_str.strip()
    if ats_system not in _WORKDAY_SYSTEMS:
        return date_str
    month_map = {
        "jan": "01", "january": "01", "feb": "02", "february": "02",
        "mar": "03", "march": "03", "apr": "04", "april": "04",
        "may": "05", "jun": "06", "june": "06", "jul": "07", "july": "07",
        "aug": "08", "august": "08", "sep": "09", "september": "09",
        "oct": "10", "october": "10", "nov": "11", "november": "11",
        "dec": "12", "december": "12",
    }
    m = re.match(r"([A-Za-z]+)\.?\s*(\d{4})\s*[-\u2013\u2014]\s*(.+)", date_str)
    if m:
        def _cvt(s: str) -> str:
            if s.lower() == "present":
                return "Present"
            mm = re.match(r"([A-Za-z]+)\.?\s*(\d{4})", s.strip())
            if mm and mm.group(1).lower() in month_map:
                return f"{month_map[mm.group(1).lower()]}/{mm.group(2)}"
            return s
        return f"{_cvt(m.group(1) + ' ' + m.group(2))} - {_cvt(m.group(3))}"
    mm = re.match(r"([A-Za-z]+)\.?\s*(\d{4})", date_str)
    if mm and mm.group(1).lower() in month_map:
        return f"{month_map[mm.group(1).lower()]}/{mm.group(2)}"
    return date_str


def _configure_base_styles(doc: Document) -> None:
    """Set the default body font + margins once up front."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    # Make sure the East-Asia font slot also points at Times so mixed
    # text doesn't fall back to Calibri on some renderers.
    rpr = style.element.get_or_add_rPr()
    for tag in ("w:rFonts",):
        existing = rpr.find(qn(tag))
        if existing is not None:
            rpr.remove(existing)
    rFonts = rpr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman",
        qn("w:hAnsi"): "Times New Roman",
        qn("w:cs"): "Times New Roman",
        qn("w:eastAsia"): "Times New Roman",
    })
    rpr.append(rFonts)

    # 0.5" margins all around
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def _add_name(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(15)
    p.paragraph_format.space_after = Pt(2)


def _add_contact(doc: Document, parts: list[str]) -> None:
    if not parts:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(" | ".join(parts))
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_section_header(doc: Document, label: str) -> None:
    """Section header = bold, underlined (via bottom border), 11pt."""
    p = doc.add_paragraph()
    run = p.add_run(label.upper())
    run.bold = True
    run.font.size = Pt(11)
    # Add a bottom border so it renders like the PDF's divider line.
    pPr = p._p.get_or_add_pPr()
    pbdr_existing = pPr.find(qn("w:pBdr"))
    if pbdr_existing is not None:
        pPr.remove(pbdr_existing)
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "auto",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def _add_entry_two_col(doc: Document, left: str, right: str, left_bold: bool = True) -> None:
    """Line with something like 'Company - Role' left and 'Dates | City' right,
    done via a right-aligned tab so the parser sees one clean line."""
    p = doc.add_paragraph()
    # Right-aligned tab at the right margin (6" with 0.5" margins + 7.5" page).
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    if left:
        run_l = p.add_run(left)
        run_l.bold = left_bold
        run_l.font.size = Pt(10.5)
    if right:
        p.add_run("\t")
        run_r = p.add_run(right)
        run_r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)


def _add_body_line(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic = italic
    p.paragraph_format.space_after = Pt(1)


def _add_bullet(doc: Document, text: str) -> None:
    # Strip any existing bullet glyphs so we only emit one clean '-'.
    text = re.sub(r"^[\u2022\u2023\u25aa\u25cf\u2013\u2014*\-]+\s*", "", text).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("- " + text)
    run.font.size = Pt(10.5)


def _render_contact(doc: Document, section: dict) -> None:
    contact = section.get("contact") or {}
    name = (contact.get("name") or "").strip()
    if name:
        _add_name(doc, name)
    parts = [
        (contact.get(f) or "").strip()
        for f in ("email", "phone", "location", "linkedin")
    ]
    parts = [p for p in parts if p]
    _add_contact(doc, parts)


def _render_education(doc: Document, section: dict, ats_system: str) -> None:
    edu = section.get("education") or {}
    if not edu:
        return
    university = (edu.get("university") or "").strip()
    location = (edu.get("location") or "").strip()
    graduation = (edu.get("graduation") or "").strip()
    if university:
        right_parts = []
        if location:
            right_parts.append(location)
        if graduation:
            right_parts.append(_format_date_for_ats(graduation, ats_system))
        _add_entry_two_col(doc, university, ", ".join(right_parts), left_bold=True)
    degree_parts = []
    major = (edu.get("major") or "").strip()
    minor = (edu.get("minor") or "").strip()
    gpa = (edu.get("gpa") or "").strip()
    if major:
        degree_parts.append(major)
    if minor:
        degree_parts.append(f"Minor: {minor}")
    if gpa:
        degree_parts.append(f"GPA: {gpa}")
    if degree_parts:
        _add_body_line(doc, " | ".join(degree_parts))
    honors = (edu.get("honors") or "").strip()
    if honors and honors.lower() != "not honors":
        _add_body_line(doc, honors)


def _render_experiences(doc: Document, section: dict, ats_system: str) -> None:
    entries = section.get("experiences") or []
    for entry in entries:
        company = (entry.get("company") or "").strip()
        role = (entry.get("role") or "").strip()
        date = (entry.get("date") or "").strip()
        location = (entry.get("location") or "").strip()
        left = ""
        if company and role:
            left = f"{company} - {role}"
        elif company:
            left = company
        elif role:
            left = role
        right_parts = []
        if date:
            right_parts.append(_format_date_for_ats(date, ats_system))
        if location:
            right_parts.append(location)
        _add_entry_two_col(doc, left, " | ".join(right_parts), left_bold=True)
        for b in entry.get("bullets") or []:
            text = b.get("text", "") if isinstance(b, dict) else str(b)
            _add_bullet(doc, text)


def _render_projects(doc: Document, section: dict, ats_system: str) -> None:
    projects = section.get("projects") or []
    for proj in projects:
        name = (proj.get("name") or "").strip()
        date = (proj.get("date") or "").strip()
        location = (proj.get("location") or "").strip()
        right_parts = []
        if date:
            right_parts.append(_format_date_for_ats(date, ats_system))
        if location:
            right_parts.append(location)
        _add_entry_two_col(doc, name, " | ".join(right_parts), left_bold=True)
        for b in proj.get("bullets") or []:
            text = b.get("text", "") if isinstance(b, dict) else str(b)
            _add_bullet(doc, text)


def _render_simple(doc: Document, section: dict) -> None:
    simple = section.get("simple") or {}
    for line in simple.get("lines") or []:
        line = line.strip() if isinstance(line, str) else str(line).strip()
        if line:
            _add_body_line(doc, line)


def build_ats_docx(sections: list[dict], ats_system: str = "greenhouse") -> bytes:
    """Build an ATS-optimized .docx from the same sections JSON as the PDF."""
    ats_system = (ats_system or "greenhouse").lower().strip()
    headers = _WORKDAY_HEADERS if ats_system in _WORKDAY_SYSTEMS else _SECTION_HEADERS

    doc = Document()
    _configure_base_styles(doc)

    for section in sections:
        if not isinstance(section, dict):
            continue
        key = (section.get("key") or "").strip()
        label = section.get("label") or ""

        if key == "contact":
            _render_contact(doc, section)
            continue

        header_label = headers.get(
            key, label.upper() if label else key.upper().replace("_", " ")
        )
        if header_label:
            _add_section_header(doc, header_label)

        if key == "education" and section.get("education"):
            _render_education(doc, section, ats_system)
        elif key in (
            "professional_experience", "research",
            "campus_involvement", "volunteer_experience",
        ) and section.get("experiences"):
            _render_experiences(doc, section, ats_system)
        elif key == "projects" and section.get("projects"):
            _render_projects(doc, section, ats_system)
        elif section.get("simple"):
            _render_simple(doc, section)
        elif section.get("experiences"):
            _render_experiences(doc, section, ats_system)
        elif section.get("projects"):
            _render_projects(doc, section, ats_system)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
