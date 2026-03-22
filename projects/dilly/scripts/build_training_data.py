#!/usr/bin/env python3
"""
Build LLM training data from your resume set.
Supports PDF (with optional OCR for image/blurry PDFs) and DOCX.

Runs the rule-based Meridian auditor on every file in RESUME_DIR and saves
(resume excerpt, scores, findings) as few-shot examples for the LLM.

Usage (from workspace root):
  python -m projects.dilly.scripts.build_training_data
  RESUME_DIR=path/to/resumes python -m projects.dilly.scripts.build_training_data

Optional env:
  MERIDIAN_PDF_OCR_THRESHOLD=100  — below this many chars from pypdf we try OCR (default 100). OCR needs: pip install pdf2image pytesseract; system: poppler, tesseract-ocr.

Output: projects/meridian/prompts/training_data.json
"""

import json
import os
import sys

# Resolve workspace root (parent of projects/)
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECTS_DIR = os.path.dirname(os.path.dirname(SCRIPT_DIR))
WORKSPACE_ROOT = os.path.dirname(PROJECTS_DIR)
if WORKSPACE_ROOT not in sys.path:
    sys.path.insert(0, WORKSPACE_ROOT)

RESUME_DIR = os.environ.get("RESUME_DIR", os.path.join(WORKSPACE_ROOT, "assets", "resumes"))
OUTPUT_PATH = os.environ.get("MERIDIAN_TRAINING_DATA")
if not OUTPUT_PATH:
    OUTPUT_PATH = os.path.join(WORKSPACE_ROOT, "projects", "dilly", "prompts", "training_data.json")
OUTPUT_PATH = os.path.abspath(OUTPUT_PATH)
# Per-example resume length cap for few-shot (keep prompt size manageable)
EXCERPT_CHARS = int(os.environ.get("MERIDIAN_EXCERPT_CHARS", "2400"))
# Min chars from native PDF extraction below which we try OCR (when pdf2image + pytesseract available)
PDF_OCR_THRESHOLD = int(os.environ.get("MERIDIAN_PDF_OCR_THRESHOLD", "100"))


def _person_key_from_filename(filename: str) -> str:
    """Derive a dedupe key from filename when candidate_name is Unknown (e.g. 'michael zeltser')."""
    import re
    base = os.path.splitext(filename)[0]
    # Strip second extension for .docx.pdf etc.
    base = os.path.splitext(base)[0] if base.lower().endswith(".docx") else base
    base = base.replace("_", " ")
    base = re.sub(r"\s*\(\d+\)\s*$", "", base)  # remove (2), (1), etc.
    base = re.sub(r"\b(resume|résumé|cv)\b", "", base, flags=re.IGNORECASE)
    base = re.sub(r"\s+", " ", base).strip().lower()
    return base[:60] if base else filename.lower()


def _pdf_text_via_pypdf(pdf_path: str) -> str:
    """Extract text using pypdf only."""
    try:
        import pypdf
        reader = pypdf.PdfReader(pdf_path)
        return "\n".join([p.extract_text() or "" for p in reader.pages]).strip()
    except Exception:
        return ""


def _pdf_text_via_ocr(pdf_path: str) -> str:
    """Extract text via OCR (pdf2image + pytesseract). Requires poppler and tesseract installed."""
    try:
        import pdf2image
        import pytesseract
    except ImportError as e:
        return ""  # caller will see empty and keep pypdf result
    try:
        images = pdf2image.convert_from_path(pdf_path, dpi=200)
        parts = []
        for img in images:
            text = pytesseract.image_to_string(img)
            if text and text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts).strip() if parts else ""
    except Exception:
        return ""


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF. Uses pypdf first; if too short, tries OCR (image/blurry PDFs) when deps available."""
    text = _pdf_text_via_pypdf(pdf_path)
    if len((text or "").strip()) < PDF_OCR_THRESHOLD:
        ocr_text = _pdf_text_via_ocr(pdf_path)
        if ocr_text and len(ocr_text.strip()) > len((text or "").strip()):
            return ocr_text
    return text or ""


def extract_text_from_docx(docx_path: str) -> str:
    """Extract plain text from a .docx file (paragraphs + tables). Requires python-docx."""
    try:
        from docx import Document
        doc = Document(docx_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_text(path: str, filename: str) -> str:
    """Dispatch to PDF or DOCX extractor by extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if lower.endswith(".docx"):
        return extract_text_from_docx(path)
    return ""


def main():
    from dilly_core.auditor import run_audit

    if not os.path.isdir(RESUME_DIR):
        print(f"Resume directory not found: {RESUME_DIR}")
        print("Set RESUME_DIR to a folder containing your PDF/DOCX resumes, or add files to assets/resumes.")
        sys.exit(1)

    allowed = (".pdf", ".docx")
    files = [f for f in os.listdir(RESUME_DIR) if f.lower().endswith(allowed)]
    if not files:
        print(f"No .pdf or .docx files found in {RESUME_DIR}")
        sys.exit(1)

    print(f"Auditing {len(files)} resumes in {RESUME_DIR} with rule-based engine...")
    examples = []
    seen_names = set()  # dedupe: one training example per person (by name from resume)

    for filename in sorted(files):
        path = os.path.join(RESUME_DIR, filename)
        text = _extract_text(path, filename)
        if not text or len(text) < 50:
            print(f"  Skip {filename}: no text extracted (for .docx install python-docx; for image PDFs install pdf2image + pytesseract and system poppler + tesseract)")
            continue
        try:
            result = run_audit(text, candidate_name="Unknown", major="Unknown", gpa=None, filename=filename)
        except Exception as e:
            print(f"  Skip {filename}: {e}")
            continue
        # One entry per person: skip if we already have this candidate
        name_key = result.candidate_name.strip().lower() if result.candidate_name else ""
        if name_key and name_key != "unknown":
            if name_key in seen_names:
                print(f"  Skip {filename}: duplicate ({result.candidate_name})")
                continue
            seen_names.add(name_key)
        else:
            # When name is Unknown, dedupe by filename-derived person key (e.g. same person, two PDFs)
            file_key = "file:" + _person_key_from_filename(filename)
            if file_key in seen_names:
                print(f"  Skip {filename}: duplicate (same file person)")
                continue
            seen_names.add(file_key)

        excerpt = text[:EXCERPT_CHARS] + ("..." if len(text) > EXCERPT_CHARS else "")
        examples.append({
            "filename": filename,
            "resume_excerpt": excerpt,
            "candidate_name": result.candidate_name,
            "major": result.major,
            "track": result.track,
            "smart_score": result.smart_score,
            "grit_score": result.grit_score,
            "build_score": result.build_score,
            "final_score": result.final_score,
            "audit_findings": result.audit_findings,
            "evidence_smart": " ".join(result.evidence_smart) if result.evidence_smart else "",
            "evidence_grit": " ".join(result.evidence_grit) if result.evidence_grit else "",
        })
        print(f"  {filename} -> {result.candidate_name} | {result.track} S:{result.smart_score:.0f} G:{result.grit_score:.0f} B:{result.build_score:.0f}")

    if not examples:
        print("WARNING: 0 examples produced. Not overwriting existing training_data.json (no text extracted from files).")
        print("Install python-docx for .docx; for image PDFs install pdf2image + pytesseract and system poppler + tesseract.")
        return

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    with open(OUTPUT_PATH, "w") as f:
        json.dump({"examples": examples, "count": len(examples)}, f, indent=2)
    print(f"Saved {len(examples)} examples to {OUTPUT_PATH}")
    print("The LLM auditor will use these as few-shot examples when MERIDIAN_FEW_SHOT=1 (or default).")


if __name__ == "__main__":
    main()
