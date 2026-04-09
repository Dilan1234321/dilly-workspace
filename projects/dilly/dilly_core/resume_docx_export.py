"""
Resume DOCX export — ATS-parseable Microsoft Word output.

Produces a `.docx` file that every major ATS (Workday, Taleo, iCIMS,
Greenhouse, Lever, Ashby, SuccessFactors) can parse without loss. The
golden rules for a parseable DOCX:

    - No text boxes, tables, columns, or shapes.
    - One linear stream of paragraphs.
    - Standard section headings Workday / Taleo look for verbatim.
    - Native paragraph bullets (real list paragraphs, not "• " strings).
    - Calibri / Arial at 10-11pt. No custom fonts.
    - Hyperlinks as plain text (parsers strip <w:hyperlink> anyway).

Three templates mirror the PDF renderer:

    tech     — compact, Skills near top, GitHub link prominent
    business — conservative spacing, centered contact header
    academic — Education / Research / Publications front-loaded

Entry point:

    render_resume_docx(sections, template='tech') -> bytes
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Template definitions ────────────────────────────────────────────────────

TEMPLATES: Dict[str, Dict[str, Any]] = {
    "tech": {
        "label": "Tech",
        "accent_rgb": (30, 64, 175),  # blue-800
        "heading_size": 11,
        "body_size": 10,
        "section_order": ["contact", "skills", "experience", "projects", "education"],
        "contact_center": False,
        "margin_inches": 0.55,
    },
    "business": {
        "label": "Business",
        "accent_rgb": (17, 24, 39),   # near-black
        "heading_size": 11,
        "body_size": 10.5,
        "section_order": ["contact", "education", "experience", "projects", "skills"],
        "contact_center": True,
        "margin_inches": 0.7,
    },
    "academic": {
        "label": "Academic",
        "accent_rgb": (31, 41, 55),
        "heading_size": 11,
        "body_size": 10.5,
        "section_order": ["contact", "education", "research", "experience", "projects", "skills"],
        "contact_center": False,
        "margin_inches": 0.65,
    },
}


# ── Helpers ────────────────────────────────────────────────────────────────

def _set_margins(doc: Document, inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_name(doc: Document, text: str, template: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template["contact_center"] else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True


def _add_contact_line(doc: Document, parts: List[str], template: Dict[str, Any]) -> None:
    if not parts:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if template["contact_center"] else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("  ·  ".join(parts))
    run.font.size = Pt(template["body_size"])
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_section_heading(doc: Document, label: str, template: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label.upper())
    run.font.size = Pt(template["heading_size"])
    run.font.bold = True
    r, g, b = template["accent_rgb"]
    run.font.color.rgb = RGBColor(r, g, b)
    # Underline the entire section heading — most ATS parsers pick up
    # the visible divider without needing a table row
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "cccccc")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_body_line(doc: Document, text: str, template: Dict[str, Any],
                    bold: bool = False, space_after: float = 1) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(template["body_size"])
    run.font.bold = bold


def _add_bullet(doc: Document, text: str, template: Dict[str, Any]) -> None:
    if not text or not text.strip():
        return
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.18)
    run = p.add_run(text)
    run.font.size = Pt(template["body_size"])


# ── Section renderers ─────────────────────────────────────────────────────

def _render_contact(doc: Document, section: Dict[str, Any], template: Dict[str, Any]) -> None:
    c = section.get("contact") or {}
    if c.get("name"):
        _add_name(doc, str(c["name"]), template)
    line_parts: List[str] = []
    for key in ("email", "phone", "location", "linkedin"):
        val = c.get(key)
        if val:
            line_parts.append(str(val))
    _add_contact_line(doc, line_parts, template)


def _render_education(doc: Document, section: Dict[str, Any], template: Dict[str, Any]) -> None:
    e = section.get("education") or {}
    if not any((e.get("university"), e.get("major"), e.get("graduation"))):
        return
    _add_section_heading(doc, "Education", template)
    line1_parts: List[str] = []
    if e.get("university"): line1_parts.append(str(e["university"]))
    if e.get("location"):   line1_parts.append(str(e["location"]))
    if line1_parts:
        _add_body_line(doc, " — ".join(line1_parts), template, bold=True)
    line2_parts: List[str] = []
    if e.get("major"):      line2_parts.append(str(e["major"]))
    if e.get("minor"):      line2_parts.append(f"Minor: {e['minor']}")
    if e.get("graduation"): line2_parts.append(str(e["graduation"]))
    if e.get("gpa"):        line2_parts.append(f"GPA: {e['gpa']}")
    if e.get("honors"):     line2_parts.append(str(e["honors"]))
    if line2_parts:
        _add_body_line(doc, " · ".join(line2_parts), template, space_after=4)


def _render_experiences(doc: Document, section: Dict[str, Any], template: Dict[str, Any],
                         heading_label: str = "Experience") -> None:
    experiences = section.get("experiences") or []
    if not experiences:
        return
    _add_section_heading(doc, heading_label, template)
    for exp in experiences:
        header_parts: List[str] = []
        company = exp.get("company", "")
        role    = exp.get("role", "")
        if role and company:
            header_parts.append(f"{company} — {role}")
        elif company:
            header_parts.append(company)
        elif role:
            header_parts.append(role)
        date_parts: List[str] = []
        if exp.get("date"):     date_parts.append(str(exp["date"]))
        if exp.get("location"): date_parts.append(str(exp["location"]))
        header_line = header_parts[0] if header_parts else ""
        if date_parts:
            header_line = f"{header_line}  ({' · '.join(date_parts)})" if header_line else f"({' · '.join(date_parts)})"
        if header_line:
            _add_body_line(doc, header_line, template, bold=True)
        for b in exp.get("bullets") or []:
            text = (b or {}).get("text") or ""
            _add_bullet(doc, text, template)


def _render_projects(doc: Document, section: Dict[str, Any], template: Dict[str, Any]) -> None:
    projects = section.get("projects") or []
    if not projects:
        return
    _add_section_heading(doc, "Projects", template)
    for p in projects:
        name = p.get("name", "")
        date = p.get("date", "")
        header = name if name else ""
        if date:
            header = f"{header}  ({date})" if header else f"({date})"
        if header:
            _add_body_line(doc, header, template, bold=True)
        for b in p.get("bullets") or []:
            text = (b or {}).get("text") or ""
            _add_bullet(doc, text, template)


def _render_simple(doc: Document, section: Dict[str, Any], template: Dict[str, Any],
                    heading_label: Optional[str] = None) -> None:
    simple = section.get("simple") or {}
    lines = [str(l or "").strip() for l in (simple.get("lines") or []) if l]
    lines = [l for l in lines if l]
    if not lines:
        return
    label = heading_label or section.get("label") or section.get("key", "").replace("_", " ").title() or "Section"
    _add_section_heading(doc, label, template)
    for line in lines:
        _add_body_line(doc, line, template)


# ── Main render ───────────────────────────────────────────────────────────

_SECTION_RENDERERS_BY_KEY = {
    "contact":                lambda d, s, t: _render_contact(d, s, t),
    "education":              lambda d, s, t: _render_education(d, s, t),
    "experience":             lambda d, s, t: _render_experiences(d, s, t, "Experience"),
    "professional_experience": lambda d, s, t: _render_experiences(d, s, t, "Experience"),
    "projects":               lambda d, s, t: _render_projects(d, s, t),
    "skills":                 lambda d, s, t: _render_simple(d, s, t, "Skills"),
}


def render_resume_docx(sections: List[Dict[str, Any]],
                        template_name: str = "tech") -> bytes:
    """Render a resume DOCX from a list of section dicts. Returns bytes."""
    if template_name not in TEMPLATES:
        template_name = "tech"
    template = TEMPLATES[template_name]

    doc = Document()
    _set_margins(doc, template["margin_inches"])

    # Set default font on the Normal style so runs that don't override
    # inherit a universally parseable face.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(template["body_size"])

    by_key: Dict[str, Dict[str, Any]] = {}
    for s in sections or []:
        if isinstance(s, dict) and s.get("key"):
            by_key[str(s["key"])] = s

    rendered: set = set()
    for wanted in template["section_order"]:
        section = by_key.get(wanted)
        if not section and wanted == "experience":
            section = by_key.get("professional_experience")
        if not section:
            continue
        renderer = _SECTION_RENDERERS_BY_KEY.get(wanted) or _SECTION_RENDERERS_BY_KEY.get(
            "experience" if "experience" in wanted else wanted
        )
        if not renderer:
            continue
        try:
            renderer(doc, section, template)
            rendered.add(section.get("key") or wanted)
        except Exception:
            continue

    # Render any sections not in the template_order at the bottom (e.g.
    # summary, honors_awards, certifications, coursework) as simple
    # lines sections so user data isn't silently dropped.
    for key, section in by_key.items():
        if key in rendered:
            continue
        if key == "professional_experience" and "professional_experience" in rendered:
            continue
        renderer = _SECTION_RENDERERS_BY_KEY.get(key)
        try:
            if renderer:
                renderer(doc, section, template)
            elif section.get("simple"):
                _render_simple(doc, section, template)
        except Exception:
            continue

    buf = BytesIO()
    doc.save(buf)
    out = buf.getvalue()
    buf.close()
    return out


__all__ = ["render_resume_docx", "TEMPLATES"]
