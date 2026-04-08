#!/usr/bin/env python3
"""
golden_resumes_test.py — Compare the legacy auditor to the new rubric_scorer
against a panel of real resumes. Prints a formatted table AND writes a JSON
dump of every result so the founder can eyeball calibration decisions.

## Usage

    cd /Users/dilankochhar/.openclaw/workspace/projects/dilly
    python3 tests/golden_resumes_test.py

## What it does

1. Finds all PDF/DOCX resumes in dilly-workspace/assets/resumes/ (or accepts
   a directory argument).
2. For each resume, extracts text and runs the existing dilly_core.auditor
   run_audit() to get the legacy AuditorResult (flat fields).
3. For each resume, also runs dilly_core.rubric_scorer.score_for_cohorts()
   against EVERY cohort rubric currently defined in
   knowledge/cohort_rubrics.json.
4. Prints a side-by-side comparison table.
5. Writes the full raw output to tests/output/golden_resumes_<timestamp>.json
   so you can diff between runs or load it into a notebook.

## Why the comparison matters

This is the "does the new scoring feel closer to true" test. The founder
explicitly said cohort_scoring_weights.py is bullshit and wants it replaced.
The comparison table lets you eyeball:

  - Does the legacy auditor produce scores that feel random or consistent?
  - Does the rubric scorer produce scores that differentiate cohorts for
    the same resume in plausible ways? (e.g., a finance major with a DCF
    should score higher in Finance than in Data Science)
  - Do low scores feel defensible or demoralizing?

Nothing in this script modifies any production code. It's read-only on the
code and write-only to tests/output/.
"""

from __future__ import annotations

import json
import os
import sys
import time
from typing import Any, Dict, List, Optional, Tuple

# Ensure we can import dilly_core regardless of CWD
_HERE = os.path.dirname(os.path.abspath(__file__))
_DILLY_ROOT = os.path.normpath(os.path.join(_HERE, ".."))
_WORKSPACE_ROOT = os.path.normpath(os.path.join(_DILLY_ROOT, "..", ".."))
for p in (_DILLY_ROOT, _WORKSPACE_ROOT):
    if p not in sys.path:
        sys.path.insert(0, p)

from dilly_core.auditor import run_audit
from dilly_core.resume_parser import parse_resume
from dilly_core.scoring import extract_scoring_signals
from dilly_core.rubric_scorer import (
    load_rubrics,
    score_with_rubric,
    RubricScore,
)


# ---------------------------------------------------------------------------
# File discovery and text extraction
# ---------------------------------------------------------------------------

_DEFAULT_RESUME_DIR = os.path.normpath(
    os.path.join(_DILLY_ROOT, "dilly-workspace", "assets", "resumes")
)


def discover_resumes(directory: Optional[str] = None) -> List[str]:
    """Return a list of absolute paths to PDF/DOCX resumes in the directory."""
    directory = directory or _DEFAULT_RESUME_DIR
    if not os.path.isdir(directory):
        return []
    files: List[str] = []
    for name in sorted(os.listdir(directory)):
        low = name.lower()
        if low.endswith(".pdf") or low.endswith(".docx"):
            files.append(os.path.join(directory, name))
    return files


def extract_text(path: str) -> str:
    """Extract raw text from a PDF or DOCX. Returns empty string on failure."""
    low = path.lower()
    if low.endswith(".pdf"):
        return _extract_pdf(path)
    if low.endswith(".docx"):
        return _extract_docx(path)
    return ""


def _extract_pdf(path: str) -> str:
    try:
        import pypdf  # noqa: F401
        reader = pypdf.PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
    except Exception as exc:
        sys.stderr.write(f"[extract] PDF failed for {path}: {exc}\n")
        return ""


def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        sys.stderr.write(f"[extract] DOCX failed for {path}: {exc}\n")
        return ""


# ---------------------------------------------------------------------------
# Per-resume scoring
# ---------------------------------------------------------------------------

def score_one_resume(
    path: str,
    rubrics: Dict[str, Dict[str, Any]],
) -> Dict[str, Any]:
    """
    Run both the legacy auditor and the rubric scorer against a single resume.
    Returns a dict describing all outcomes.
    """
    filename = os.path.basename(path)
    record: Dict[str, Any] = {
        "filename": filename,
        "extraction_ok": False,
        "parsed": None,
        "legacy": None,
        "rubric_scores": {},
        "errors": [],
    }

    # 1. Extract text
    text = extract_text(path)
    if not text or len(text) < 50:
        record["errors"].append(f"text_extraction_empty ({len(text)} chars)")
        return record
    record["extraction_ok"] = True

    # 2. Parse
    try:
        parsed = parse_resume(text, filename=filename)
        record["parsed"] = {
            "name": parsed.name,
            "major": parsed.major,
            "gpa": parsed.gpa,
            "section_keys": list((parsed.sections or {}).keys())[:10],
            "text_length": len(text),
        }
    except Exception as exc:
        record["errors"].append(f"parse_resume: {exc}")
        return record

    normalized_text = parsed.normalized_text or text

    # 3. Legacy auditor (existing path — untouched by this session)
    try:
        legacy = run_audit(
            normalized_text,
            candidate_name=parsed.name or "Unknown",
            major=parsed.major or "Unknown",
            gpa=parsed.gpa,
            filename=filename,
        )
        record["legacy"] = {
            "track": legacy.track,
            "smart": round(legacy.smart_score, 1),
            "grit":  round(legacy.grit_score,  1),
            "build": round(legacy.build_score, 1),
            "final": round(legacy.final_score, 1),
        }
    except Exception as exc:
        record["errors"].append(f"run_audit: {exc}")

    # 4. Rubric scorer — against every cohort rubric
    try:
        signals = extract_scoring_signals(normalized_text, gpa=parsed.gpa, major=parsed.major or "")
        record["signals_summary"] = {
            "gpa": signals.gpa,
            "work_entries": signals.work_entry_count,
            "impacts": signals.quantifiable_impact_count,
            "leadership": signals.outcome_leadership_count,
            "has_research": signals.has_research,
            "deployed_app": signals.deployed_app_or_live_link,
            "recognized_tech_employer": signals.recognized_tech_employer,
            "honors": signals.honors_count,
        }
        for cid, rubric in rubrics.items():
            try:
                result = score_with_rubric(signals, normalized_text, rubric)
                record["rubric_scores"][cid] = {
                    "smart": round(result.smart, 1),
                    "grit":  round(result.grit, 1),
                    "build": round(result.build, 1),
                    "composite": round(result.composite, 1),
                    "recruiter_bar": result.recruiter_bar,
                    "above_bar": result.above_bar,
                    "matched_count": len(result.matched_signals),
                    "unmatched_count": len(result.unmatched_signals),
                    "unmeasured_count": len(result.unmeasured_signals),
                    "top_matched": [s.signal for s in result.matched_signals[:5]],
                    "top_unmatched": [s.signal for s in result.unmatched_signals[:5]],
                }
            except Exception as exc:
                record["errors"].append(f"score_with_rubric({cid}): {exc}")
    except Exception as exc:
        record["errors"].append(f"extract_scoring_signals: {exc}")

    return record


# ---------------------------------------------------------------------------
# Pretty printing
# ---------------------------------------------------------------------------

def _fmt(n: Optional[float], width: int = 5) -> str:
    if n is None:
        return "  —  "
    return f"{n:>{width}.1f}"


def print_summary(records: List[Dict[str, Any]]) -> None:
    """Print a human-readable comparison table to stdout."""
    print("=" * 100)
    print("GOLDEN RESUMES TEST — Legacy auditor vs new rubric scorer")
    print("=" * 100)
    print()

    n_ok = sum(1 for r in records if r.get("legacy") and r.get("rubric_scores"))
    n_total = len(records)
    print(f"Scored: {n_ok}/{n_total} resumes")
    print()

    # For each resume, print legacy and all rubric scores
    for rec in records:
        fn = rec["filename"][:60]
        print(f"── {fn}")
        if rec.get("errors"):
            print(f"   errors: {rec['errors'][:3]}")
        parsed = rec.get("parsed") or {}
        print(f"   parsed: name={parsed.get('name')}  major={parsed.get('major')}  gpa={parsed.get('gpa')}")
        sigs = rec.get("signals_summary") or {}
        print(f"   signals: work={sigs.get('work_entries')}  impacts={sigs.get('impacts')}  lead={sigs.get('leadership')}  research={sigs.get('has_research')}  deployed={sigs.get('deployed_app')}  honors={sigs.get('honors')}")
        legacy = rec.get("legacy") or {}
        if legacy:
            print(f"   LEGACY [{legacy.get('track'):>10}]  "
                  f"S:{_fmt(legacy.get('smart'))}  "
                  f"G:{_fmt(legacy.get('grit'))}  "
                  f"B:{_fmt(legacy.get('build'))}  "
                  f"= {_fmt(legacy.get('final'))}")
        rubric_scores = rec.get("rubric_scores") or {}
        # Sort by composite descending so the best-fit cohort appears first
        sorted_rubrics = sorted(
            rubric_scores.items(),
            key=lambda kv: kv[1].get("composite", 0),
            reverse=True,
        )
        for cid, scores in sorted_rubrics:
            bar = scores.get("recruiter_bar", 0)
            mark = "✓" if scores.get("above_bar") else "·"
            print(f"   {mark} RUBRIC [{cid:28}]  "
                  f"S:{_fmt(scores.get('smart'))}  "
                  f"G:{_fmt(scores.get('grit'))}  "
                  f"B:{_fmt(scores.get('build'))}  "
                  f"= {_fmt(scores.get('composite'))}  "
                  f"(bar {bar:.0f})")
        print()

    print("=" * 100)
    print("LEGEND:")
    print("  ✓ = composite score is at or above the cohort recruiter bar")
    print("  · = below the bar")
    print("  S / G / B = Smart / Grit / Build (0-100)")
    print("  Legacy uses dilly_core/auditor.py (untouched). Rubric uses new rubric_scorer.py + cohort_rubrics.json.")
    print("=" * 100)


# ---------------------------------------------------------------------------
# Output persistence
# ---------------------------------------------------------------------------

def write_output(records: List[Dict[str, Any]], rubrics: Dict[str, Dict[str, Any]]) -> str:
    """Write the full JSON dump to tests/output/<timestamp>.json."""
    out_dir = os.path.join(_HERE, "output")
    os.makedirs(out_dir, exist_ok=True)
    ts = time.strftime("%Y%m%d_%H%M%S")
    path = os.path.join(out_dir, f"golden_resumes_{ts}.json")
    payload = {
        "generated_at": ts,
        "rubrics_tested": sorted(rubrics.keys()),
        "rubric_count": len(rubrics),
        "resume_count": len(records),
        "results": records,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2, default=str)
    return path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    resume_dir = sys.argv[1] if len(sys.argv) > 1 else _DEFAULT_RESUME_DIR
    resumes = discover_resumes(resume_dir)
    if not resumes:
        print(f"No resumes found in {resume_dir}")
        return 1

    # Cap to first 8 for performance; override with env var
    cap = int(os.environ.get("DILLY_GOLDEN_CAP", "8"))
    resumes = resumes[:cap]

    try:
        rubrics = load_rubrics()
    except Exception as exc:
        print(f"Failed to load rubrics: {exc}")
        return 2

    print(f"Loaded {len(rubrics)} rubrics: {sorted(rubrics.keys())}")
    print(f"Scoring {len(resumes)} resumes...\n")

    records: List[Dict[str, Any]] = []
    for path in resumes:
        t0 = time.time()
        rec = score_one_resume(path, rubrics)
        rec["elapsed_sec"] = round(time.time() - t0, 2)
        records.append(rec)
        fn = os.path.basename(path)[:50]
        status = "OK" if rec.get("rubric_scores") else "FAIL"
        print(f"  [{status}] {fn} ({rec['elapsed_sec']:.1f}s)")

    print()
    print_summary(records)

    out_path = write_output(records, rubrics)
    print(f"\nFull JSON output saved to: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
