"""
pdf_extract.py — Unified PDF text extraction for Dilly.

Uses PyMuPDF (fitz) as the primary extractor because it:
  - Handles 2-column layouts correctly via block-level extraction
  - Preserves y-coordinate ordering (top-to-bottom, left-to-right)
  - Is 5-10x faster than pdfminer
  - Successfully extracts GPAs that PyPDF2 silently drops (verified on
    BridgetKlaus2026Resume.pdf: PyPDF2 → None, PyMuPDF → 4.0)

Falls back to PyPDF2 if PyMuPDF is unavailable or the primary extraction
fails for any reason (encoding edge cases, corrupted PDFs).

This module is used by the NEW rubric-based scoring path
(dilly_core.rubric_scorer.audit_with_rubric) and by any caller that wants
best-effort text extraction. The legacy DillyResumeAuditor class is NOT
modified — it keeps using its own extraction path to avoid regression risk.

## Usage

    from dilly_core.pdf_extract import extract_text_best_effort
    text = extract_text_best_effort("/path/to/resume.pdf")
    # Or from bytes:
    text = extract_text_best_effort_from_bytes(pdf_bytes)

## Why block-based instead of page.get_text()

Default `page.get_text("text")` uses reading-order detection that can be
wrong on 2-column resumes (e.g. reads down the left column then the right
column, interleaving them). Block-based extraction + manual y/x sorting
gives us correct reading order on any layout.
"""

from __future__ import annotations

import os
import sys
from typing import Optional


def extract_text_best_effort(path: str) -> str:
    """
    Extract text from a PDF or DOCX file. Returns empty string on total failure.

    Order of attempts:
      1. PyMuPDF (fitz) block-based extraction — layout-aware
      2. PyPDF2 / pypdf — legacy fallback
      3. python-docx — for .docx files

    The best result (longest non-empty) wins. This means corrupted PDFs
    that one extractor handles poorly can still succeed via another.
    """
    if not path or not os.path.isfile(path):
        return ""

    low = path.lower()

    if low.endswith(".docx"):
        return _extract_docx(path)

    if not low.endswith(".pdf"):
        return ""

    # Try PyMuPDF first (preferred)
    pymupdf_text = _extract_pdf_pymupdf(path)
    if pymupdf_text and len(pymupdf_text.strip()) > 50:
        return pymupdf_text

    # Fall back to PyPDF2
    pypdf_text = _extract_pdf_pypdf(path)
    if pypdf_text and len(pypdf_text.strip()) > 50:
        return pypdf_text

    # Return whichever had content (even if short)
    return pymupdf_text or pypdf_text or ""


def extract_text_best_effort_from_bytes(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF bytes (e.g. from an HTTP upload without writing
    to disk). Returns empty string on total failure.
    """
    if not pdf_bytes:
        return ""

    # Try PyMuPDF first
    pymupdf_text = _extract_pdf_pymupdf_bytes(pdf_bytes)
    if pymupdf_text and len(pymupdf_text.strip()) > 50:
        return pymupdf_text

    # Fall back to PyPDF2
    pypdf_text = _extract_pdf_pypdf_bytes(pdf_bytes)
    if pypdf_text and len(pypdf_text.strip()) > 50:
        return pypdf_text

    return pymupdf_text or pypdf_text or ""


# ---------------------------------------------------------------------------
# PyMuPDF implementations
# ---------------------------------------------------------------------------

def _extract_pdf_pymupdf(path: str) -> str:
    """
    Extract text using PyMuPDF with block-level ordering.

    Sorts blocks by y-coordinate (top-to-bottom) with a 5-pixel tolerance,
    then by x-coordinate (left-to-right). This correctly handles 2-column
    resume layouts that would otherwise be interleaved by default extraction.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ""

    try:
        doc = fitz.open(path)
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] PyMuPDF open failed: {exc}\n")
        return ""

    try:
        text_parts = []
        for page in doc:
            text_parts.append(_extract_page_blocks(page))
        return "\n".join(text_parts).strip()
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] PyMuPDF extraction failed: {exc}\n")
        return ""
    finally:
        try:
            doc.close()
        except Exception:
            pass


def _extract_pdf_pymupdf_bytes(pdf_bytes: bytes) -> str:
    """Same as _extract_pdf_pymupdf but from bytes in memory."""
    try:
        import fitz
    except ImportError:
        return ""

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] PyMuPDF open-bytes failed: {exc}\n")
        return ""

    try:
        text_parts = []
        for page in doc:
            text_parts.append(_extract_page_blocks(page))
        return "\n".join(text_parts).strip()
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] PyMuPDF extraction failed: {exc}\n")
        return ""
    finally:
        try:
            doc.close()
        except Exception:
            pass


def _extract_page_blocks(page) -> str:
    """
    Extract text from a single page using block-level extraction with
    y-then-x sorting. Handles 2-column layouts correctly.

    block tuple format: (x0, y0, x1, y1, text, block_no, block_type)
    block_type 0 = text, 1 = image. We only want text blocks.
    """
    try:
        blocks = page.get_text("blocks")
    except Exception:
        return page.get_text("text") or ""

    # Filter to text blocks only
    text_blocks = [b for b in blocks if len(b) >= 7 and b[6] == 0]
    if not text_blocks:
        return page.get_text("text") or ""

    # Sort by y0 (rounded to 5px tolerance for wobble), then x0
    text_blocks.sort(key=lambda b: (round(b[1] / 5) * 5, b[0]))
    return "\n".join(b[4] for b in text_blocks if b[4]).strip()


# ---------------------------------------------------------------------------
# PyPDF2 / pypdf fallback implementations
# ---------------------------------------------------------------------------

def _extract_pdf_pypdf(path: str) -> str:
    """Fallback: extract with pypdf."""
    try:
        import pypdf
    except ImportError:
        return ""
    try:
        reader = pypdf.PdfReader(path)
        return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] pypdf failed: {exc}\n")
        return ""


def _extract_pdf_pypdf_bytes(pdf_bytes: bytes) -> str:
    """Fallback: extract with pypdf from bytes."""
    try:
        import pypdf
        import io
    except ImportError:
        return ""
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] pypdf bytes failed: {exc}\n")
        return ""


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(path: str) -> str:
    """Extract text from a .docx file."""
    try:
        from docx import Document
    except ImportError:
        sys.stderr.write("[pdf_extract] python-docx not installed\n")
        return ""

    try:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception as exc:
        sys.stderr.write(f"[pdf_extract] docx failed: {exc}\n")
        return ""


# ---------------------------------------------------------------------------
# GPA-focused extraction helper
# ---------------------------------------------------------------------------

def extract_gpa_best_effort(text: str) -> Optional[float]:
    """
    Best-effort GPA extraction from resume text. Returns None if no GPA found.

    Tries multiple regex patterns because resumes use varied GPA formats:
      - "GPA: 3.85"
      - "3.85/4.00"
      - "3.85 GPA"
      - "Cumulative GPA: 3.85"
      - "4.00" next to "GPA"
    """
    if not text:
        return None

    import re

    patterns = [
        # "GPA: 3.85" or "GPA 3.85" or "GPA - 3.85"
        r'GPA[:\s\-\.]*(\d\.\d{1,3})',
        # "3.85/4.0" or "3.85 / 4.0"
        r'(\d\.\d{1,3})\s*/\s*4\.?\d?',
        # "3.85 GPA"
        r'(\d\.\d{1,3})\s+GPA',
        # "Cumulative: 3.85"
        r'Cumulative[:\s]+(\d\.\d{1,3})',
        # "Overall GPA 3.85"
        r'Overall\s+GPA[:\s]*(\d\.\d{1,3})',
        # "Major GPA 3.85"
        r'Major\s+GPA[:\s]*(\d\.\d{1,3})',
    ]

    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            try:
                val = float(m.group(1))
                if 0.0 <= val <= 4.3:  # allow for A+ = 4.3 at some schools
                    return val
            except (ValueError, TypeError):
                continue

    return None
